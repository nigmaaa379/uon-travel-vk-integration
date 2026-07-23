import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_borders(table, border_color="000000", border_size="4"):
    tblPr = table._tbl.tblPr
    borders_elm = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:left w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:bottom w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:right w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:insideH w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:insideV w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders_elm)

def build_proposal_doc():
    doc = docx.Document()
    
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        
        # Footer
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("Стр. ")
        f_run.font.name = "Times New Roman"
        f_run.font.size = Pt(9)
        fldSimple = parse_xml(r'<w:fldSimple %s w:instr="PAGE"/>' % nsdecls('w'))
        f_p._p.append(fldSimple)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11.5)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(15)

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(16)
        r = p.add_run(text)
        r.font.italic = True
        r.font.size = Pt(11)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(13)

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(12)

    def add_p(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
        r = p.add_run(text)
        r.font.italic = italic
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
        p.add_run(text)
        return p

    def add_box(text, title=""):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.rows[0].cells[0]
        cell.width = Inches(6.7)
        set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
        set_cell_shading(cell, "F4F4F4")
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="none"/>\n'
            f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="000000"/>\n'
            f'  <w:bottom w:val="none"/>\n'
            f'  <w:right w:val="none"/>\n'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        if title:
            rb = p.add_run(title + "\n")
            rb.font.bold = True
            rb.font.size = Pt(10.5)
        r = p.add_run(text)
        r.font.size = Pt(10)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # CONTENT
    add_title("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ:\nРАЗРАБОТКА ВЕБ-САЙТА, БЛОГА И ИНТЕГРАЦИЙ С QUI-QUO И CRM U-ON TRAVEL")
    add_subtitle("Проект для турагентства «Семейный Атлас» (@turiotfamilyatlas)\nДата составления: 23 июля 2026 г.")

    add_box(
        "Настоящее предложение разработано для турагентства «Семейный Атлас». Проект направлен на создание современного, удобного веб-сайта на базе WordPress, подчеркивающего философию премиального семейного отдыха («Бережём отдых и нервы туристов»), с интеграцией подборок Qui-Quo, спецблока «Отели дня», полноценного Блога и автопередачей заявок в U-ON Travel.",
        "РЕЗЮМЕ ПРОЕКТА"
    )

    # Section 1
    add_h1("1. КОНЦЕПЦИЯ И ВИЗУАЛЬНОЕ ПОЗИЦИОНИРОВАНИЕ")
    add_bullet("Философия бренда: «Семейный отдых без стресса. Бережём отдых и нервы туристов». Агентство берет на себя все заботы по подбору и сопровождению семей с детьми.", "Ключевая идея: ")
    add_bullet("Фирменный стиль: Разработка адаптивного дизайна на базе зарегистрированного товарного знака и логотипа компании, а также стилистики Telegram-канала t.me/turiotfamilyatlas.", "Визуал и логотип: ")
    add_bullet("Эстетика: Тёплые, премиальные тона, вызывающие чувство надёжности, уюта и комфорта. Удобный интерфейс как для мобильных устройств, так и для ПК.", "Интерфейс: ")

    # Section 2
    add_h1("2. СОСТАВ И ФУНКЦИОНАЛ ВЕБ-САЙТА (WORDPRESS)")
    add_p("Сайт разрабатывается на платформе WordPress (удобное управление контентом) с оптимизированными самописными модулями для максимальной скорости загрузки.")

    add_h2("2.1. Структура страниц и обязательные разделы:")
    add_bullet("Главная страница — Презентация бренда, ключевые преимущества семейного отдыха, быстрый подбор туров, акценты продаж и отзывы клиентов.", "• ")
    add_bullet("Раздел «Отели дня» — Свободно обновляемая витрина проверенных семейных отелей с фото, рейтингом и детальным описанием условий для детей.", "• ")
    add_bullet("Раздел «Блог» — Полноценный журнал со статьями, лайфхаками по поездкам с детьми, обзорами стран и SEO-оптимизацией для привлечения органического трафика.", "• ")
    add_bullet("Разделы популярных стран — Страницы по ключевым направлениям (Турция, ОАЭ, Мальдивы, Сейшелы, Таиланд и др.).", "• ")
    add_bullet("Юридический блок (Роспотребнадзор) — Информация об ИП/ООО, ИНН, ОГРН, реестровый номер ЕФРТА, контакты, часы работы, Публичная оферта и правила возврата.", "• ")
    add_bullet("Защита персональных данных (РКН / 152-ФЗ) — Чекбоксы согласий без предпроставленных галочек, Политика конфиденциальности, Cookie Consent всплывающий баннер.", "• ")

    add_box(
        "Примечание по структуре: По согласованию из структуры исключены лишние разделы («Круизы» и «Экскурсии»), что сделает сайт максимально сфокусированным на семейных турах и отелях.",
        "ИСКЛЮЧЕНИЯ ИЗ ТЗ"
    )

    # Section 3
    add_h1("3. ИНТЕГРАЦИИ (QUI-QUO, CRM U-ON TRAVEL И МЕССЕНДЖЕРЫ)")

    add_h2("3.1. Интеграция с сервисом подборок Qui-Quo (Кви-Кво):")
    add_bullet("Внедрение фирменных подборок и виджетов Qui-Quo на страницы сайта и в статьи блога.", "• ")
    add_bullet("Возможность мгновенно просмотреть тур из подборки Кви-Кво и оставить заявку на бронирование.", "• ")

    add_h2("3.2. CRM-система U-ON Travel:")
    add_bullet("Все заявки с форм обратной связи, «Отелей дня» и подборок автоматически отправляются в U-ON Travel в виде карточки «Обращение».", "• ")
    add_bullet("Автозаполнение параметров тура (страна, отель, состав семьи, бюджет, контакты).", "• ")

    add_h2("3.3. Уведомления и альтернатива Личному Кабинету (ЛК):")
    add_bullet("Альтернатива сложному ЛК: Вместо громоздкой регистрации на сайте, турист авторизуется и получает сопровождение напрямую через мессенджеры (Telegram / VK).", "Решение: ")
    add_bullet("Авто-уведомление руководителя на Email и в мессенджер при поступлении новой заявки.", "Уведомления: ")

    # Section 4 - Schedule & Cost Table
    add_h1("4. ЭТАПЫ РАБОТ И СТОИМОСТЬ")

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, "000000", "6")

    headers = ["Этап / Наименование работ", "Описание и результат", "Срок"]
    widths = [Inches(2.2), Inches(3.5), Inches(1.0)]

    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        set_cell_shading(cell, "CCCCCC")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h_text)
        r.font.bold = True
        r.font.size = Pt(9.5)

    data = [
        ("Этап 1. Проектирование и Визуал", "Сбор логотипа/стиля ВК/ТГ, верстка структуры WordPress, адаптивный дизайн для мобильных устройств.", "3-4 дня"),
        ("Этап 2. Блог, Отели дня и Контент", "Настройка модуля Блога, витрины «Отели дня», страниц стран, юридического блока Роспотребнадзора и 152-ФЗ.", "4-5 дней"),
        ("Этап 3. Интеграции Qui-Quo & U-ON", "Подключение виджетов подборок Кви-Кво, автосоздание обращений в U-ON Travel, настройка уведомлений.", "3-4 дня"),
        ("Этап 4. SEO и Финальный запуск", "Настройка sitemap, robots.txt, мета-тегов, подключения Яндекс Метрики/Карт, тестирование и передача прав.", "2 дня")
    ]

    for r_idx, r_data in enumerate(data):
        row = table.add_row()
        for c_idx, val in enumerate(r_data):
            cell = row.cells[c_idx]
            cell.width = widths[c_idx]
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F9F9F9")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(val)
            r.font.size = Pt(9)
            if c_idx == 0:
                r.font.bold = True
            if c_idx == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_box(
        "• Итоговый срок разработки: 12-14 рабочих дней.\n"
        "• Все юридические требования Роспотребнадзора и Роскомнадзора учитываются «из коробки».\n"
        "• Интеграция с U-ON Travel и Qui-Quo обеспечивает автоматизацию продаж без лишней ручной работы.",
        "ГАРАНТИИ И ИТОГ"
    )

    desktop_path = "/Users/ape.ces/Desktop/КП_Сайт_Семейный_Атлас_QuiQuo_UON.docx"
    doc.save(desktop_path)
    print(f"Proposal saved to: {desktop_path}")

if __name__ == "__main__":
    build_proposal_doc()
