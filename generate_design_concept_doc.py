import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_borders(table, border_color="000000", border_size="4"):
    tblPr = table._tbl.tblPr
    borders_elm = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:left w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:bottom w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:right w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:insideH w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:insideV w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders_elm)

def create_design_doc():
    doc = docx.Document()
    
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        
        # Footer
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("Стр. ")
        f_run.font.name = "Times New Roman"
        f_run.font.size = Pt(9)
        fldSimple = parse_xml(r'<w:fldSimple %s w:instr="PAGE"/>' % nsdecls('w'))
        f_p._p.append(fldSimple)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11.5)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(15)

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(16)
        r = p.add_run(text)
        r.font.italic = True
        r.font.size = Pt(11)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(13)

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(12)

    def add_p(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
        r = p.add_run(text)
        r.font.italic = italic
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
        p.add_run(text)
        return p

    def add_box(text, title=""):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.rows[0].cells[0]
        cell.width = Inches(6.7)
        set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
        set_cell_shading(cell, "F4F4F4")
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="none"/>\n'
            f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="000000"/>\n'
            f'  <w:bottom w:val="none"/>\n'
            f'  <w:right w:val="none"/>\n'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        if title:
            rb = p.add_run(title + "\n")
            rb.font.bold = True
            rb.font.size = Pt(10.5)
        r = p.add_run(text)
        r.font.size = Pt(10)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # DOCUMENT CONTENT
    add_title("ДИЗАЙН-КОНЦЕПЦИЯ И АРХИТЕКТУРА САЙТА:\nТУРАГЕНТСТВО «СЕМЕЙНЫЙ АТЛАС»")
    add_subtitle("Презентация стилевого решения, типографики и структуры страниц\n(на базе зарегистрированного логотипа и Telegram-канала t.me/turiotfamilyatlas)\nДата: 23 июля 2026 г.")

    add_box(
        "Настоящий документ представляет визуальное решение веб-сайта для турагентства «Семейный Атлас». Концепция опирается на философию спокойного, надежного и комфортного семейного отдыха («Бережём ваш отдых и нервы»). В дизайне сочетаются мягкие тёплые оттенки, высокая читаемость шрифтов, удобные блоки «Отели дня», Блог и подборки Qui-Quo.",
        "КОНЦЕПЦИЯ БРЕНДА"
    )

    add_h1("1. ЦВЕТОВАЯ ПАЛИТРА И СТИЛЕВАЯ ГАРМОНИЯ")
    add_p("Дизайн выдержан в эстетике премиального семейного клуба без кричащих дешевых демпинговых цветов:")
    add_bullet("Глубокий морской/океанический тотемный цвет (#0F4C5C). Символизирует стабильность, надежность и премиальное качество обслуживания.", "Основной цвет брендбука: ")
    add_bullet("Тёплый солнечный терракот / янтарный (#E36414). Используется для кнопок целевого действия (CTA), акцентов и важных уведомлений.", "Акцентный цвет: ")
    add_bullet("Мягкий песочно-кремовый (#FAF8F5). Обеспечивает уютное восприятие и легкое чтение контента с любых экранов.", "Фоновый цвет: ")
    add_bullet("Тёмно-графитовый (#1D2D35). Текст читается мягко и не утомляет глаза.", "Цвет текста: ")

    add_h1("2. ТИПОГРАФИКА И ШРИФТОВАЯ ПАРА")
    add_bullet("Заголовки и акценты: Современный элегантный геометрик Outfit / Playfair Display. Придает страницам солидность и характер премиального журнала.", "1. ")
    add_bullet("Основной текст: Высокочитаемый гротеск Plus Jakarta Sans / Inter. Оптимизирован для мобильных устройств, обеспечивает четкое чтение сведений о турах и отелях.", "2. ")

    add_h2("3. СТРУКТУРА СТРАНИЦ И СМЫСЛОВЫЕ БЛОКИ")

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, "000000", "6")

    headers = ["Блок сайта", "Смысловое наполнение и Визуал", "Задачи блока"]
    widths = [Inches(1.8), Inches(3.7), Inches(1.2)]

    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        set_cell_shading(cell, "CCCCCC")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h_text)
        r.font.bold = True
        r.font.size = Pt(9.5)

    data = [
        ("Шапка и Главный баннер (Hero)", "Заголовок «Бережём ваш семейный отдых и нервы», три плашки преимуществ (детские клубы, проверенный трансфер, поддержка 24/7) + форма экспресс-подбора.", "Первое впечатление, позиционирование, доверие."),
        ("Витрина «Отели дня»", "Карточки проверенных семейных отелей с отметками для детей (аквапарки, меню, Rixy Club), ценой и кнопкой «Запросить».", "Быстрые продажи рекомендованных отелей."),
        ("Подборки Qui-Quo", "Интерактивный виджет с подборками туров от экспертов компании с прямым переходом к бронированию.", "Удобство изучения вариантов туристами."),
        ("Раздел «Блог родителям»", "Статьи с советами педиатров, гайдами по перелетам с малышами, обзорами пляжей с песчаным входом.", "Завоевание экспертности и SEO-трафик."),
        ("Отзывы и Доверие", "Интеграция официального бейджа и отзывов Яндекс Карт / Яндекс Бизнес с оценкой 5.0.", "Подтверждение реальной репутации."),
        ("Юридический подвал", "Официальные данные ИП/ООО, реестровый номер ЕФРТА, оферта, политика ПД (152-ФЗ) и дисклеймер.", "Соблюдение законов РФ (РКН и Роспотребнадзор).")
    ]

    for r_idx, r_data in enumerate(data):
        row = table.add_row()
        for c_idx, val in enumerate(r_data):
            cell = row.cells[c_idx]
            cell.width = widths[c_idx]
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F9F9F9")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(val)
            r.font.size = Pt(9)
            if c_idx == 0:
                r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_box(
        "На Рабочем столе сохранен интерактивный файл концепта Дизайн_концепт_Семейный_Атлас.html. Вы можете открыть его в любом браузере, чтобы визуально оценить шрифты, карточки «Отелей дня» и оформление блоков на ПК и смартфонах.",
        "ИНТЕРАКТИВНЫЙ ПРОТОТИП"
    )

    desktop_path = "/Users/ape.ces/Desktop/Дизайн_концепция_и_структура_Семейный_Атлас.docx"
    doc.save(desktop_path)
    print(f"Design document saved to: {desktop_path}")

if __name__ == "__main__":
    create_design_doc()
