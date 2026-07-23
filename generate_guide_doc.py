import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_borders(table, border_color="000000", border_size="4"):
    tblPr = table._tbl.tblPr
    borders_elm = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:left w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:bottom w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:right w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:insideH w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:insideV w:val="single" w:sz="{border_size}" w:space="0" w:color="{border_color}"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders_elm)

def create_guide_document():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        
        # Footer Page Number
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("Стр. ")
        f_run.font.name = "Times New Roman"
        f_run.font.size = Pt(9)
        f_run.font.color.rgb = RGBColor(0, 0, 0)
        
        fldSimple = parse_xml(r'<w:fldSimple %s w:instr="PAGE"/>' % nsdecls('w'))
        f_p._p.append(fldSimple)

    # Base Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11.5)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(16)
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(16)
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11.5)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0, 0, 0)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
        r.font.bold = True

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.bold = True

    def add_p(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
        r = p.add_run(text)
        r.font.italic = italic
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
        p.add_run(text)
        return p

    def add_box(text, title=""):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.rows[0].cells[0]
        cell.width = Inches(6.7)
        set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
        set_cell_shading(cell, "F2F2F2")
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="none"/>\n'
            f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="000000"/>\n'
            f'  <w:bottom w:val="none"/>\n'
            f'  <w:right w:val="none"/>\n'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        if title:
            rb = p.add_run(title + "\n")
            rb.font.bold = True
            rb.font.size = Pt(10.5)
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.italic = True
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # DOCUMENT CONTENT
    add_title("ПОШАГОВАЯ ИНСТРУКЦИЯ ДЛЯ ЗАКАЗЧИКА:\nГДЕ НАЙТИ КЛЮЧИ, ТОКЕНЫ И РЕКВИЗИТЫ ДЛЯ ЗАПУСКА САЙТА И ИНТЕГРАЦИЙ")
    add_subtitle("Инструкция по получению всех необходимых данных для турагентства «Сбежим на море» / «Tur Exotic»\nДата формирования: 23 июля 2026 г.")

    add_box(
        "Данная инструкция содержит пошаговый порядок действий для сбора ключей доступа, API-токенов, ссылок и юридических реквизитов. Все указанные данные необходимы разработчику для полной интеграции сайта с CRM U-ON Travel, чат-ботами, модулем Турвизор и сервисами Яндекса.",
        "НАЗНАЧЕНИЕ ДОКУМЕНТА"
    )

    # 1. Legal Requirements
    add_h1("1. ЮРИДИЧЕСКИЕ ДАННЫЕ И РЕКВИЗИТЫ (ТРЕБОВАНИЯ РОСПОТРЕБНАДЗОРА И РКН)")
    add_p("Данная информация публикуется в подвале (футере) сайта и на странице «Контакты» для соблюдения Закона «О защите прав потребителей» и ФЗ № 132-ФЗ «Об основах туристской деятельности в РФ»:")
    
    table1 = doc.add_table(rows=1, cols=3)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.autofit = False
    set_table_borders(table1, "000000", "6")

    headers1 = ["Наименование поля", "Где посмотреть / Откуда взять", "Пример заполнения"]
    widths1 = [Inches(1.8), Inches(2.6), Inches(2.3)]
    
    for i, h_text in enumerate(headers1):
        cell = table1.rows[0].cells[i]
        cell.width = widths1[i]
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        set_cell_shading(cell, "CCCCCC")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h_text)
        r.font.bold = True
        r.font.size = Pt(9.5)

    data1 = [
        ("Полное наименование", "Свидетельство ИНН/ОГРН или выписка из ЕГРИП / ЕГРЮЛ", "ИП Баляева Ольга Юрьевна / ООО «Туристическая компания»"),
        ("ИНН / ОГРНИП / ОГРН", "Свидетельство о постановке на учет в налоговом органе", "ИНН 650106422378 / ОГРНИП 318650100012345"),
        ("Номер в ЕФРТА (реестр турагентов)", "Личный кабинет на сайте реестра турагентов (tourism.gov.ru / ЕФРТА)", "РТА 0000000 (или номер из уведомления о внесении в реестр)"),
        ("Адрес компании", "Договор аренды офиса / юридический адрес", "344000, г. Ростов-на-Дону, пр-кт Стачки, 59, офис 218"),
        ("Контакты и режим", "Официальный рабочий телефон и e-mail турагентства", "+7 (909) 400-96-19, Пн-Пт 10:00–19:00, e-mail: hello@example.ru")
    ]

    for r_idx, r_data in enumerate(data1):
        row = table1.add_row()
        for c_idx, val in enumerate(r_data):
            cell = row.cells[c_idx]
            cell.width = widths1[c_idx]
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F9F9F9")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(val)
            r.font.size = Pt(9)
            if c_idx == 0:
                r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 2. Tourvisor
    add_h1("2. ТУРВИЗОР (TOURVISOR) — ПОИСКОВЫЙ МОДУЛЬ И API ГОРЯЩИХ ТУРОВ")
    add_p("Модуль Турвизора отвечает за выгрузку актуальных цен на туры, форму поиска на сайте и рассылку горящих туров в чат-ботах.")

    add_h2("2.1. Как найти API Token (Ключ доступа API Турвизор):")
    add_bullet("Зайдите на официальный сайт tourvisor.ru и войдите в личный кабинет турагентства.", "Шаг 1: ")
    add_bullet("В верхнем меню откройте раздел «Настройки» → далее вкладку «API» (или «Интеграция с CRM»).", "Шаг 2: ")
    add_bullet("В поле «Токен API» скопируйте длинную строку из букв и цифр (например: tv_token_123456789abc).", "Шаг 3: ")
    add_bullet("Примечание: Если раздел API неактивен, напишите в техподдержку Турвизора через форму на сайте: «Просим активировать доступ к XML/API для подключения сайта и интеграции».", "Шаг 4: ")

    add_h2("2.2. Как получить код виджета поиска туров для сайта:")
    add_bullet("В личном кабинете Турвизора перейдите в раздел «Модули» → «Поиск туров» (или «Мои модули»).", "Шаг 1: ")
    add_bullet("Выберите или настройте внешний вид модуля поиска под цветовые гаммы вашего сайта.", "Шаг 2: ")
    add_bullet("Нажмите кнопку «Получить код» и скопируйте весь блок кода (содержит теги <script src=\"//tourvisor.ru/module/init.js\"></script>).", "Шаг 3: ")

    # 3. U-ON CRM
    add_h1("3. CRM-СИСТЕМА U-ON TRAVEL")
    add_p("U-ON Travel служит центром обработки обращений: все заявки с сайта, из ВК и из чат-ботов попадают в раздел «Обращения».")

    add_h2("Как найти API-ключ U-ON:")
    add_bullet("Войдите в личный кабинет U-ON Travel под администратором.", "Шаг 1: ")
    add_bullet("Перейдите в меню «Настройки» → далее «Интеграции» → «API / Webhooks».", "Шаг 2: ")
    add_bullet("Включите переключатель «Разрешить работу по API» и скопируйте ваш API-ключ (например: 5Ao0v48kP4HgZOWBe...).", "Шаг 3: ")
    add_bullet("В разделе «Настройки» → «Сотрудники» уточните ID менеджера (по умолчанию ID: 3), на которого должны автоматически назначаться заявки.", "Шаг 4: ")

    # 4. VK
    add_h1("4. СООБЩЕСТВО ВКОНТАКТЕ (VK CALLBACK API & BOT)")
    add_p("Необходимо для работы автоквалификации туриста в ВК и отправки уведомлений руководителю.")

    add_h2("Как получить Ключ доступа и ID группы ВКонтакте:")
    add_bullet("Откройте вашу группу ВКонтакте под правами Администратора.", "Шаг 1: ")
    add_bullet("Перейдите в «Управление» → «Работа с API».", "Шаг 2: ")
    add_bullet("Нажмите «Создать ключ», отметьте галочками пункты «Сообщения сообщества» и «Доступ к справочной информации», затем нажмите «Сохранить» и скопируйте Ключ доступа (начинается на vk1.a...).", "Шаг 3: ")
    add_bullet("В адресной строке или настройках группы скопируйте цифровой ID группы (например: 228564119).", "Шаг 4: ")
    add_bullet("Укажите свой личный VK ID (цифровой номер вашей личной страницы ВКонтакте), на который должны приходить сообщения о новых заявках.", "Шаг 5: ")

    # 5. Yandex
    add_h1("5. СЕРВИСЫ ЯНДЕКСА (ЯНДЕКС МЕТРИКА И ЯНДЕКС БИЗНЕС)")

    add_h2("5.1. Яндекс Метрика (Номер счётчика аналитики):")
    add_bullet("Перейдите на сайт metrika.yandex.ru под вашим логином Яндекс.", "Шаг 1: ")
    add_bullet("В списке счетчиков найдите ваш сайт — под названим укажите 8-значный номер счетчика (например: 98765432).", "Шаг 2: ")
    add_bullet("Если счетчика нет — нажмите желтую кнопку «Добавить счетчик», введите адрес сайта и нажмите «Создать». Скопируйте номер.", "Шаг 3: ")

    add_h2("5.2. Яндекс Бизнес / Карты (Виджет отзывов и рейтинга):")
    add_bullet("Перейдите в кабинет b2b.yandex.ru (Яндекс Бизнес) или найдите карточку вашей компании на Карты Яндекс (yandex.ru/maps).", "Шаг 1: ")
    add_bullet("Скопируйте прямую ссылку на карточку компании (например: https://yandex.ru/maps/org/tur_exotic/123456789).", "Шаг 2: ")
    add_bullet("В кабинете Яндекс Бизнес откройте меню «О компании» → «Промоматериалы» → скопируйте код виджета рейтинга для сайта.", "Шаг 3: ")

    # 6. SMTP
    add_h1("6. КОРПОРАТИВНАЯ ПОЧТА ДЛЯ УВЕДОМЛЕНИЙ (SMTP)")
    add_p("Чтобы сайт мог отправлять письма с подтверждениями и уведомления о заявках менеджерам, понадобятся данные корпоративной почты:")
    add_bullet("E-mail адрес рассылки (например: hello@example.ru или notify@example.ru).", "1. ")
    add_bullet("Пароль для внешних приложений (в Mail.ru / Yandex делается в Безопасность → Пароли для внешних приложений).", "2. ")
    add_bullet("E-mail адрес(-а) руководителей и менеджеров, на которые должны приходить письма о заявках.", "3. ")

    # Output save
    desktop_path = "/Users/ape.ces/Desktop/Инструкция_для_заказчика_данные_и_ключи.docx"
    doc.save(desktop_path)
    print(f"File successfully created: {desktop_path}")

if __name__ == "__main__":
    create_guide_document()
